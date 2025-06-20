import os
import cv2
from PyPDF2 import PdfReader
import numpy as np
from pdf2image import convert_from_path
import pytesseract
from docx import Document

def extract_text_pdf(file_path):
    reader = PdfReader(file_path)
    page_texts = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        page_texts.append((f"page_{i + 1}", text))
    return page_texts

def extract_text_docx(file_path):
    doc = Document(file_path)
    texts = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            texts.append((f"para_{i + 1}", para.text.strip()))
    return texts

def extract_text_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return [("full", text.strip())]

def perform_ocr_pdf(file_path):
    pages = convert_from_path(file_path)
    ocr_texts = []
    for i, page in enumerate(pages):
        text = pytesseract.image_to_string(page)
        ocr_texts.append((f"ocr_page_{i + 1}", text.strip()))
    return ocr_texts

def perform_ocr_docx(file_path):
    doc = Document(file_path)
    ocr_texts = []
    img_index = 0

    for rel in doc.part._rels.values():
        if "image" in rel.target_ref:
            img_index += 1
            img_data = rel.target_part.blob
            img_array = np.frombuffer(img_data, np.uint8)
            image = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

            text = pytesseract.image_to_string(image)
            ocr_texts.append((f"ocr_img_{img_index}", text.strip()))
    
    return ocr_texts

def load_document(file_path):
    """Load document and return list of (source, text)"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_pdf(file_path) + perform_ocr_pdf(file_path)
    elif ext == ".docx":
        return extract_text_docx(file_path) + perform_ocr_docx(file_path)
    elif ext == ".txt":
        return extract_text_txt(file_path)
    else:
        raise ValueError("Unsupported file type.")
